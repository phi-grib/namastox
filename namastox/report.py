#! -*- coding: utf-8 -*-

# Description    NAMASTOX command
#
# Authors:       Manuel Pastor (manuel.pastor@upf.edu)
#
# Copyright 2022 Manuel Pastor
#
# This file is part of NAMASTOX
#
# NAMASTOX is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation version 3.
#
# Flame is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with NAMASTOX. If not, see <http://www.gnu.org/licenses/>.

from namastox.logger import get_logger
from namastox.ra import Ra
import os
import yaml
from rdkit import Chem
from rdkit.Chem import Draw
import xlsxwriter
from datetime import date
from namastox.utils import id_generator
import docx
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# from docx import Document
# from docx.shared import Pt
# from docx.shared import RGBColor

LOG = get_logger(__name__)

def generateSubstanceImage (smiles, oname):
    try:
        m = Chem.MolFromSmiles(smiles)
    except:
        return False

    Draw.MolToFile(m, oname, imageType='png')
    return True

def report_excel (ra):
    reportfile = os.path.join (ra.rapath,'report.xlsx')

    if os.path.isfile(reportfile):
        try:
            os.remove(reportfile)
        except:
            return False, 'Failed! the report file is open or not writtable'
        
    workbook = xlsxwriter.Workbook(reportfile)
    worksheet = workbook.add_worksheet()

    # define styles and formats
    worksheet.set_column(0,0,width=40)
    worksheet.set_column(1,2,width=25)
    worksheet.set_column(3,3,width=60)
    worksheet.set_column(4,9,width=25)

    label_format = workbook.add_format({'align':'top', 'text_wrap': True, 'bold': True})
    value_format = workbook.add_format({'align':'top', 'text_wrap': True, 'font_color':'#3465A4'})

    # General section
    raitem = ra.general
    labels = ['title','endpoint', 'general_description', 'administration_route', 'regulatory_framework','species','background']

    irow =0 
    worksheet.write(irow, 0, 'General Information', label_format )
    for ilabel in labels:
        if not ilabel in raitem:
            continue
        worksheet.write(irow, 1, ilabel.replace('_',' '), label_format )
        worksheet.write(irow, 3, raitem[ilabel], value_format )
        irow+=1

    substances_items = ra.general['substances']

    # generate substance image
    if 'smiles' in substances_items:
        generateSubstanceImage(ra, substances_items['smiles'])

    substance_keys = ['name', 'casrn', 'id', 'smiles']
    
    for i,isubstance in enumerate(substances_items):
        worksheet.write(irow, 1, 'substance', label_format ) 

        # generate substance image
        if 'smiles' in isubstance:
            spath = os.path.join(ra.rapath, f'substance-{str(i)}.png')
            if os.path.isfile(spath):
                success = True
            else:
                success = generateSubstanceImage(isubstance['smiles'],spath)
            if success:
                worksheet.set_row(irow, 60)
                worksheet.write(irow, 2, 'structure', label_format )
                worksheet.insert_image (irow, 3, spath, {"x_scale": 0.25, "y_scale": 0.25})
                irow+=1

        for ikey in substance_keys:
            if not ikey in isubstance:
                continue
            worksheet.write(irow, 2, ikey, label_format )
            worksheet.write(irow, 3, isubstance[ikey], value_format )
            irow+=1

    irow+=1

    # Results section

    bool_to_text = {True:'Yes', False:'No'}
    workflow = ra.workflow

    for reitem in ra.results:

        # Name and Description are not in results but in workflow
        inode = workflow.getNode(reitem['id'])
        itask = inode.getTask()
        idict = itask.getDescriptionDict()
        name = idict['task description']['name']
        description = idict['task description']['description']
        if 'label' in reitem:
            label = reitem['label']
        else:
            label = ''

        worksheet.write(irow, 0, name+f" ({label})", label_format )

        worksheet.write(irow, 1, 'description', label_format )
        worksheet.write(irow, 3, description, value_format )
        irow+=1
        
        worksheet.write(irow, 1, 'summary', label_format )
        worksheet.write(irow, 3, reitem['summary'], value_format )
        irow+=1

        if 'decision' in reitem:
            worksheet.write(irow, 1, 'decision', label_format )
            worksheet.write(irow, 3, bool_to_text[reitem['decision']], value_format )
            irow+=1

            worksheet.write(irow, 1, 'justification', label_format )
            worksheet.write(irow, 3, reitem['justification'], value_format )
            irow+=1
        
        else:
            if reitem['result_type'] == 'text':

                for iresult in reitem['values']:
                    worksheet.write(irow, 1, 'result', label_format )
                    worksheet.write(irow, 3, iresult, value_format )
                    irow+=1
                
                if 'uncertainties' in reitem:
                    for iresult in reitem['uncertainties']:
                        if 'uncertainty' in iresult and iresult['uncertainty'] != '':
                            worksheet.write(irow, 1, 'uncertainty', label_format )
                            worksheet.write(irow, 3, iresult['uncertainty'], value_format )
                            irow+=1
            
                        if 'term' in iresult and iresult['term'] != '':
                            worksheet.write(irow, 1, 'term', label_format )
                            worksheet.write(irow, 3, iresult['term'], value_format )
                            irow+=1


            elif reitem['result_type'] == 'value':

                if 'methods' in reitem:
                    if len(reitem['methods'])>0:
                        worksheet.write(irow, 1, 'method', label_format )

                        for iresult in reitem['methods']:
                            worksheet.write(irow, 2, iresult['name'], label_format )

                            if 'description' in iresult:
                                worksheet.write(irow, 3, iresult['description'], value_format )

                            if 'link' in iresult:
                                worksheet.write(irow, 4, iresult['link'], value_format )

                            if 'sensitivity' in iresult:
                                worksheet.write(irow, 5, f"sens.:{iresult['sensitivity']}", value_format )

                            if 'specificity' in iresult:
                                worksheet.write(irow, 6, f"spec.:{iresult['specificity']}", value_format )

                            if 'sd' in iresult:
                                worksheet.write(irow, 7, f"sd:{iresult['sd']}", value_format )

                            irow+=1
                        
                if len(reitem['values'])==len(reitem['uncertainties']):
                    worksheet.write(irow, 1, 'result', label_format )
                    
                    for iresult,iuncertain in zip(reitem['values'],reitem['uncertainties']):
                            
                        if 'substance' in iresult:
                            worksheet.write(irow, 2, iresult['substance'], label_format )

                        if 'parameter' in iresult:
                            worksheet.write(irow, 3, iresult['parameter'], value_format )
                            
                        if 'value' in iresult:
                            worksheet.write(irow, 4, iresult['value'], value_format )

                        if 'unit' in iresult and iresult['unit']!='':
                            worksheet.write(irow, 5, iresult['unit'], value_format )
                        
                        if 'uncertainty' in iuncertain and iuncertain['uncertainty'] != '':
                            worksheet.write(irow, 6, iuncertain['uncertainty'], value_format )
            
                        if 'term' in iuncertain and iuncertain['term'] != '':
                            worksheet.write(irow, 7, iuncertain['term'], value_format )

                        if 'method' in iresult and iresult['method'] != '':
                            worksheet.write(irow, 8, f"method: {iresult['method']}", value_format )
                            
                        irow+=1

                # fallback when the size of parameters and uncertainties don't match (this should never happen)
                else:
                    worksheet.write(irow, 1, 'result', label_format )
                    
                    for iresult in reitem['values']:
                            
                        if 'substance' in iresult:
                            worksheet.write(irow, 2, iresult['substance'], label_format )

                        if 'parameter' in iresult:
                            worksheet.write(irow, 3, iresult['parameter'], label_format )
                            
                        if 'value' in iresult:
                            worksheet.write(irow, 4, iresult['value'], value_format )

                        if 'unit' in iresult and iresult['unit']!='':
                            worksheet.write(irow, 5, iresult['unit'], value_format )
                    
                        irow+=1
                            
                    for iuncertain in reitem['uncertainties']:
                        if 'uncertainty' in iuncertain and iuncertain['uncertainty'] != '':
                            worksheet.write(irow, 2, 'uncertainty', value_format )
                            worksheet.write(irow, 3, iuncertain['uncertainty'], value_format )
                            irow+=1
            
                        if 'term' in iuncertain and iuncertain['term'] != '':
                            worksheet.write(irow, 2, 'term', value_format )
                            worksheet.write(irow, 3, iuncertain['term'], value_format )
                            irow+=1

    
        if len(reitem['links'])> 0:    
            worksheet.write(irow, 1, 'links', label_format )
            for ilink in reitem['links']:
                if 'include' in ilink and not ilink['include']:
                    continue 
                worksheet.write(irow, 2, ilink['label'].replace('_',' '), label_format )
                worksheet.write(irow, 3, ilink['File'], value_format )
                irow+=1

        worksheet.write(irow, 1, 'date', label_format )
        worksheet.write(irow, 3, reitem['date'], value_format )
        irow+=1

    irow+=1

    # Notes section
    for noitem in ra.notes:
        worksheet.write(irow, 0, f"{noitem['title']} ({noitem['id']})", label_format )
        worksheet.write(irow, 1, 'note', label_format )
        worksheet.write(irow, 3, noitem['text'], value_format )
        irow+=1
        worksheet.write(irow, 1, 'date', label_format )
        worksheet.write(irow, 3, noitem['date'], value_format )
        irow+=1
    
        irow+=1

    irow+=1
        
    # close workbook    
    workbook.close()

    return True, reportfile

def addGeneralSection (document, ra, item):
    if item in ra.general and ra.general[item]!=None and len(ra.general[item])>2:
        item_title = item.capitalize().replace('_',' ')
        document.add_heading(item_title, level=3)
        document.add_paragraph (ra.general[item])

def addHyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def insertText (cell, val):
    cell.text = str(val) 

def addResult (document, ra, reitem, section, order):
    bool_to_text = {True:'Yes', False:'No'}
    workflow = ra.workflow

    # Name and Description are not in results but in workflow
    inode = workflow.getNode(reitem['id'])
    itask = inode.getTask()
    idict = itask.getDescriptionDict()
    name = idict['task description']['name']
    description = idict['task description']['description']
    if 'label' in reitem:
        label = reitem['label']
    else:
        label = ''

    #TODO add subsection number
    document.add_heading (f"{str(section)}.{str(order)} {name} ({label})", level=2)
    t = document.add_table(rows = 1, cols = 1, style='Table Grid')
    dparagraph = t.rows[0].cells[0].paragraphs[0]
    dparagraph.add_run(description).italic = True
    
    document.add_heading ('Summary', level=3)
    document.add_paragraph(reitem['summary'])

    if 'decision' in reitem:
        document.add_heading('Decision', level=3)
        document.add_paragraph(bool_to_text[reitem['decision']])

        document.add_heading('Justification', level=3 )
        document.add_paragraph(reitem['justification'])
    
    else:
        # Table with methods
        if 'methods' in reitem:
            methods = reitem['methods']
            if len(methods)>0:

                document.add_heading('Methods', level=3 )
        
                t = document.add_table(rows = 1, cols=6)
                t.style = 'Light Grid Accent 1'
                t.autofit = True
                hdr_cells = t.rows[0].cells
                hdr_cells[0].text = 'Name'
                hdr_cells[1].text = 'Description'
                hdr_cells[2].text = 'Link'
                hdr_cells[3].text = 'Sensit.'
                hdr_cells[4].text = 'Spec.'
                hdr_cells[5].text = 'SD'
        
                for iresult in methods:
                    line_cells = t.add_row().cells
                    insertText(line_cells[0], iresult['name'])
                    if 'description' in iresult:
                        insertText(line_cells[1], iresult['description'])
                    if 'link' in iresult:
                        link = iresult['link']
                        insertText(line_cells[2], link)

                        # model documentation in included in the zip with attachements and 
                        # can be linked
                        if link.startswith('documentation_'):
                            link_p = document.add_paragraph (link+' : ', style='List Bullet')
                            addHyperlink(link_p, link, link)

                    if 'sensitivity' in iresult:
                        insertText(line_cells[3], iresult['sensitivity'])
                    if 'specificity' in iresult:
                        insertText(line_cells[4], iresult['specificity'])
                    if 'sd' in iresult:
                        insertText(line_cells[5], iresult['sd'])


        document.add_heading('Result', level=3 )

        if reitem['result_type'] == 'text':

            for iresult in reitem['values']:
                document.add_paragraph(iresult)

        elif reitem['result_type'] == 'value':
                

            # Table with values and uncertainties
            if len(reitem['values'])==len(reitem['uncertainties']):

                # check if there is at least one element for the three optional columns
                iheader = 3
                method_touch = False
                for iresult in reitem['values']:
                    if 'method' in iresult and iresult['method']!='':
                        method_touch = True
                        iheader+=1
                        break

                unit_touch = False
                for iresult in reitem['values']:
                    if 'unit' in iresult and iresult['unit']!='':
                        unit_touch = True
                        iheader+=1
                        break

                uncertainty_touch = False
                for iuncertain in reitem['uncertainties']:
                    if 'uncertainty' in iuncertain and iuncertain['uncertainty'] != 0:
                        uncertainty_touch = True
                        iheader+=1
                        break

                term_touch = False
                for iuncertain in reitem['uncertainties']:
                    if 'term' in iuncertain and iuncertain['term'] != '':
                        term_touch = True
                        iheader+=1
                        break

                t = document.add_table(rows = 1, cols=iheader)

                # t.style = 'Table Grid'
                t.style = 'Light Grid Accent 1'
                t.autofit = True
                hdr_cells = t.rows[0].cells
                hdr_cells[0].text = 'Substance'
                hdr_cells[1].text = 'Parameter'
                hdr_cells[2].text = 'Value'
                
                # add headers for optional columns
                iheader = 3
                if method_touch:
                    hdr_cells[iheader].text = 'Method'
                    iheader +=1 

                if unit_touch:
                    hdr_cells[iheader].text = 'Unit'
                    iheader +=1

                if uncertainty_touch:
                    hdr_cells[iheader].text = 'Uncertainty'
                    iheader +=1

                if term_touch:
                    hdr_cells[iheader].text = 'Term'

                for iresult,iuncertain in zip(reitem['values'],reitem['uncertainties']):
                    line_cells = t.add_row().cells
                    if 'substance' in iresult:
                        insertText(line_cells[0], iresult['substance'])

                    if 'parameter' in iresult:
                        insertText(line_cells[1], iresult['parameter'])

                    if 'value' in iresult:
                        insertText(line_cells[2], iresult['value'])

                    # for the optional columns, add only if the column exist, but even
                    # so, check if for this particular objetc we must add something
                    
                    icol = 3
                    if method_touch:
                        if 'method' in iresult and iresult['method']!='':
                            insertText(line_cells[icol], iresult['method'])
                        icol+=1
                    
                    if unit_touch:
                        if 'unit' in iresult and iresult['unit']!='':
                            insertText(line_cells[icol], iresult['unit'])
                        icol+=1

                    if uncertainty_touch:
                        if 'uncertainty' in iuncertain and iuncertain['uncertainty'] != '':
                            insertText(line_cells[icol], iuncertain['uncertainty'])
                        icol+=1

                    if term_touch:
                        if 'term' in iuncertain and iuncertain['term'] != '':
                            insertText(line_cells[icol], iuncertain['term'])
                        
    if len(reitem['links'])> 0:    
        document.add_heading('Supporting documents', level=3 )
        for ilink in reitem['links']: 
            link_p = document.add_paragraph (ilink['label'].replace('_',' ')+' : ', style='List Bullet')
            addHyperlink(link_p, ilink['File'], ilink['File'])


def report_word (ra):
    reportfile = os.path.join (ra.rapath,'report.docx')

    if os.path.isfile(reportfile):
        try:
            os.remove(reportfile)
        except:
            return False, 'Failed! the report file is open or not writtable'
        
    # path = os.path.dirname(os.path.abspath(__file__))
    # path = os.path.join(path,'default')
    # template = os.path.join(path,'generic_word.docx')
    # document = Document(template)
    document = docx.Document()

    ## define style for normal and heading 1
    # normal_style = document.styles['Normal']
    # normal_font = normal_style.font
    # normal_font.name = 'Calibri'
    # normal_font.size = Pt(10)

    # heading_style = document.styles['heading 1']
    # heading_font = heading_style.font
    # heading_font.name = 'Calibri'
    # heading_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # heading_font.size = Pt(12)

    # Title
    document.add_heading(ra.general['title'], 0)

    # Fixed sections
    document.add_paragraph (f'Report date {date.today().isoformat()}')
    document.add_paragraph ('Author')
    document.add_paragraph ('<Disclaimer>')  
    document.add_paragraph ('<How to use this report>')  

    # Table of Contents
    document.add_heading ('Table of Contents')
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "<Please right-click to update TOC>"

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    # General info section
    document.add_heading ('1. General information', level=1)

    # -> substance
    document.add_heading ('Substance', level=3)
    substances_items = ra.general['substances']

    substance_keys = ['name', 'casrn', 'id']
    
    for i,isubstance in enumerate(substances_items):

        # generate substance image
        if 'smiles' in isubstance:
            spath = os.path.join(ra.rapath, f'substance-{str(i)}.png')
            if os.path.isfile(spath):
                success = True
            else:
                success = generateSubstanceImage(isubstance['smiles'],spath)
            if success:
                document.add_picture (spath, width=Cm(4.0))

        # add name, CAS and ID's
        for ikey in substance_keys: 
            if ikey in isubstance and isubstance[ikey] != None and isubstance[ikey]!='':
                #TODO: add the structure graph, removing 'smiles' as text
                document.add_paragraph(ikey + ': '+isubstance[ikey], style='List Bullet')

        # add empty line between items
        if len(substances_items)>1:
            document.add_paragraph ('')

    # -> rest of sub-sections in General Information
    items = ['general_description', 'background', 'regulatory_framework', 'endpoint', 'species']
    for item in items:
        addGeneralSection (document, ra, item)

    # Results section

    ra_sections = [ 
        {'id':'A', 're': [], 'label': 'Preliminary'},
        {'id':'H', 're': [], 'label': 'Hazard'},
        {'id':'B', 're': [], 'label': 'ADME'},
        {'id':'E', 're': [], 'label': 'Exposure'},
        {'id':'X', 're': [], 'label': 'Integration'},
        {'id':'Z', 're': [], 'label': 'Conclusions'}
    ]

    # try to assign the results to standard results sections
    assigned = 0
    for reitem in ra.results:
        id = reitem['id'][0]

        for isection in ra_sections:
            if id == isection['id']:
                isection['re'].append(reitem)
                assigned +=1

    # if we failed to assign probably the workflow is custom 
    # so don't assign sections
    if assigned < len(ra.results):
        document.add_heading ('2. Results',level=1 )
        iorder=1
        for reitem in ra.results:
            addResult (document, ra, reitem, 2, iorder)
            iorder+=1
    else:
        isec = 1
        for isection in ra_sections:
            # skip empty sections
            if len(isection['re'])==0: 
                continue

            # assign section sequential number
            isec+=1
            document.add_heading (f'{isec}. {isection["label"]}',level=1 )
            iorder = 1
            for reitem in isection['re']:
                addResult (document, ra, reitem, isec, iorder)
                iorder+=1
    
    # Notes section. We decided to avoid numbering this section
    document.add_heading ('Notes', level=1)
    for noitem in ra.notes:
        document.add_heading (f"{noitem['title']} ({noitem['id']})", level=3)
        document.add_paragraph(noitem['text'])
    
    try:
        document.save(reportfile)
    except:
        return False, f'error saving document as {reportfile}'

    return True, reportfile


def action_report (raname, report_format):

    # instantiate a ra object
    ra = Ra(raname)
    succes, results = ra.load()

    if not succes:
        return False, results
    
    if report_format == 'yaml':
    
        reportfile = os.path.join (ra.rapath,'report.yaml')

        # include only "human interesting" sections
        dict_temp = {
            'general': ra.general, 
            'results': ra.results,
            'notes': ra.notes,
        }
        with open(reportfile,'w') as f:
            f.write(yaml.dump(dict_temp))

        return True, reportfile
    
    elif report_format == 'excel':

        return report_excel(ra)

    elif report_format == 'word':

        return report_word(ra)

    return False, 'format unsupported'

